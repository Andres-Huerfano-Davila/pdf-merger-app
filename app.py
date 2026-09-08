import io
import re
import zipfile
import hashlib
import shutil
from html import escape
import streamlit as st
import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageOps, ImageEnhance
try:
    import pytesseract
except ImportError:
    pytesseract = None
from docx import Document
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as ReportLabImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

# =========================================================
# CONFIG
# =========================================================
APP_TITLE = "📄 Suite PDF: Unir, Dividir, Convertir Word e Imágenes, Firmar y Comprimir"
TARGET_DEFAULT = "Lennin Karina Triana Fandiño"

# Aguamarina
ACCENT = "#20DE6E"
ACCENT_HOVER = "#16B85B"
OCR_AVAILABLE = pytesseract is not None and shutil.which("tesseract") is not None

st.set_page_config(page_title="Suite PDF", page_icon="📄", layout="wide")

# =========================================================
# STYLES
# =========================================================
st.markdown(
    f"""
    <style>
      .stApp {{
        background: radial-gradient(1200px 700px at 10% 10%, #F2FFFA 0%, #EFFFF7 35%, #F7FFFB 70%, #FFFFFF 100%);
      }}
      section[data-testid="stSidebar"] {{
        background: linear-gradient(180deg, #EFFFF7 0%, #E8FFF1 100%);
        border-right: 1px solid #C9F7DE;
      }}
      h1, h2, h3 {{ color: #0F5132; }}

      .hero {{
        background: linear-gradient(90deg, {ACCENT} 0%, #8CF0C0 100%);
        border: 1px solid #BDF3D4;
        padding: 18px 22px;
        border-radius: 18px;
        font-weight: 900;
        color: #063B22;
        text-align: center;
        font-size: 28px;
        box-shadow: 0 10px 30px rgba(32, 222, 110, 0.18);
        margin-bottom: 16px;
      }}

      .card {{
        background: #FFFFFF;
        border: 1px solid #C9F7DE;
        border-radius: 16px;
        padding: 16px 16px;
        box-shadow: 0 6px 18px rgba(15, 81, 50, 0.07);
        margin-bottom: 12px;
      }}

      .muted {{ color: #3D6B55; font-size: 14px; }}

      div[data-testid="stButton"] > button,
      div[data-testid="stDownloadButton"] > button {{
        background-color: {ACCENT};
        color: #05311D;
        border: none;
        border-radius: 12px;
        padding: 0.60rem 1rem;
        font-weight: 900;
      }}
      div[data-testid="stButton"] > button:hover,
      div[data-testid="stDownloadButton"] > button:hover {{
        background-color: {ACCENT_HOVER};
        color: white;
      }}

      .block-container {{
        padding-top: 1.2rem;
        padding-bottom: 2rem;
      }}
    </style>
    """,
    unsafe_allow_html=True
)

# =========================================================
# SESSION STATE
# =========================================================
def init_state():
    defaults = {
        # Merge queue
        "merge_files": [],          # [{"name":str,"bytes":bytes,"size":int}]
        "merge_signature": None,
        "merged_pdf_bytes": None,

        # Passwords (shared across modules)
        "pdf_passwords": {},        # {filename: password}
        "encrypted_files": [],      # for merge
        "encrypted_split": [],      # for split

        # Detection for signing
        "detected": False,
        "det_page": None,
        "det_rect": None,
        "det_method": None,

        # Inputs
        "last_output_name": "PDF_unido.pdf",
        "last_target": TARGET_DEFAULT,

        # Signature offsets
        "sig_dx": 0.0,
        "sig_dy": 0.0,

        # Images module
        "converted_images_pdf_bytes": [],
        "converted_images_names": [],
        "merged_images_pdf_bytes": None,

        # Compress module
        "compressed_pdf_bytes": None,
        "compressed_name": "archivo_comprimido.pdf",

        # Split module
        "split_outputs": [],        # [(filename, bytes), ...]

        # Word module
        "word_pdf_bytes": None,
        "word_pdf_name": "documento_convertido.pdf",
        "word_source_name": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

def reset_merge():
    st.session_state.merge_files = []
    st.session_state.merge_signature = None
    st.session_state.merged_pdf_bytes = None
    st.session_state.encrypted_files = []
    st.session_state.detected = False
    st.session_state.det_page = None
    st.session_state.det_rect = None
    st.session_state.det_method = None
    st.session_state.sig_dx = 0.0
    st.session_state.sig_dy = 0.0

def reset_images():
    st.session_state.converted_images_pdf_bytes = []
    st.session_state.converted_images_names = []
    st.session_state.merged_images_pdf_bytes = None

def reset_compress():
    st.session_state.compressed_pdf_bytes = None
    st.session_state.compressed_name = "archivo_comprimido.pdf"

def reset_split():
    st.session_state.split_outputs = []
    st.session_state.encrypted_split = []

def reset_word():
    st.session_state.word_pdf_bytes = None
    st.session_state.word_pdf_name = "documento_convertido.pdf"
    st.session_state.word_source_name = None

init_state()

# =========================================================
# FAST PREVIEW (lazy + cached)
# =========================================================
_PDF_BYTES_CACHE = {}  # {hash: bytes}

def _pdf_key(pdf_bytes: bytes) -> str:
    return hashlib.md5(pdf_bytes).hexdigest()

@st.cache_data(show_spinner=False, max_entries=400)
def _page_count_cached(pdf_key: str) -> int:
    b = _PDF_BYTES_CACHE[pdf_key]
    doc = fitz.open(stream=b, filetype="pdf")
    n = doc.page_count
    doc.close()
    return n

@st.cache_data(show_spinner=False, max_entries=2000)
def _render_page_cached(pdf_key: str, page_index_0: int, dpi: int) -> bytes:
    b = _PDF_BYTES_CACHE[pdf_key]
    doc = fitz.open(stream=b, filetype="pdf")
    page = doc[page_index_0]
    zoom = dpi / 72.0
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    doc.close()
    return pix.tobytes("png")

def preview_pdf_viewer_fast(pdf_bytes: bytes, key_prefix: str, title: str):
    """Visor rápido: no renderiza hasta que el usuario active el toggle."""
    if not pdf_bytes:
        return

    pdf_key = _pdf_key(pdf_bytes)
    _PDF_BYTES_CACHE[pdf_key] = pdf_bytes

    show_key = f"{key_prefix}_show"
    if show_key not in st.session_state:
        st.session_state[show_key] = False

    st.markdown(f"### 👀 {title}")
    c1, c2 = st.columns([1, 2])
    with c1:
        st.session_state[show_key] = st.toggle(
            "Mostrar previsualización",
            value=st.session_state[show_key],
            key=f"{key_prefix}_toggle"
        )
    with c2:
        st.caption("Tip: DPI 72/96 es rápido. Sube a 120/150 solo si necesitas detalle.")

    if not st.session_state[show_key]:
        st.info("Previsualización oculta para mayor velocidad.")
        return

    total_pages = _page_count_cached(pdf_key)

    dpi = st.select_slider(
        "Calidad (DPI)",
        options=[72, 96, 120, 150],
        value=96,
        key=f"{key_prefix}_dpi"
    )

    page_state = f"{key_prefix}_page"
    if page_state not in st.session_state:
        st.session_state[page_state] = 1

    nav1, nav2, nav3 = st.columns([1, 2, 1])
    with nav1:
        if st.button("⬅️", key=f"{key_prefix}_prev", disabled=st.session_state[page_state] <= 1):
            st.session_state[page_state] -= 1
            st.rerun()
    with nav3:
        if st.button("➡️", key=f"{key_prefix}_next", disabled=st.session_state[page_state] >= total_pages):
            st.session_state[page_state] += 1
            st.rerun()

    st.session_state[page_state] = st.number_input(
        "Página",
        min_value=1,
        max_value=total_pages,
        value=int(st.session_state[page_state]),
        step=1,
        key=f"{key_prefix}_page_input"
    )

    png = _render_page_cached(pdf_key, int(st.session_state[page_state]) - 1, int(dpi))
    st.image(png, use_container_width=True)

def preview_first_page_quick(pdf_bytes: bytes, key_prefix: str, label: str):
    """Preview ultra rápido (solo primera página)"""
    if not pdf_bytes:
        return
    pdf_key = _pdf_key(pdf_bytes)
    _PDF_BYTES_CACHE[pdf_key] = pdf_bytes
    st.markdown(f"#### 👁️ {label}")
    dpi = st.select_slider("DPI (rápido)", options=[72, 96, 120], value=72, key=f"{key_prefix}_dpi")
    png = _render_page_cached(pdf_key, 0, int(dpi))
    st.image(png, use_container_width=True)

# =========================================================
# CORE HELPERS
# =========================================================
def normalize(s: str) -> str:
    s = (s or "").lower().strip()
    s = re.sub(r"\s+", " ", s)
    return s

# =========================================================
# PASSWORD SUPPORT
# =========================================================
def open_pdf_with_password(pdf_bytes: bytes, filename: str):
    """
    Retorna (doc, None) o (None, "encrypted") o (None, "error: ...")
    """
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if doc.is_encrypted:
            pwd = st.session_state.pdf_passwords.get(filename, "")
            ok = doc.authenticate(pwd) if pwd is not None else False
            if not ok:
                doc.close()
                return None, "encrypted"
        return doc, None
    except Exception as e:
        return None, f"error: {type(e).__name__}: {e}"

# =========================================================
# MERGE (robust + password)
# =========================================================
def merge_pdfs_with_password(files_in_order):
    """
    Une PDFs robusto:
    - links=0 para evitar errores con links internos
    - maneja encriptados pidiendo contraseña
    returns: (merged_bytes or None, warnings_list, encrypted_list)
    """
    merged = fitz.open()
    warnings = []
    encrypted = []

    for item in files_in_order:
        name = item["name"]
        b = item["bytes"]

        doc, err = open_pdf_with_password(b, name)
        if err == "encrypted":
            encrypted.append(name)
            continue
        if doc is None:
            warnings.append(f"❌ {name}: no se pudo abrir ({err})")
            continue

        try:
            merged.insert_pdf(doc, links=0, annots=1)
        except Exception as e:
            warnings.append(f"❌ {name}: no se pudo unir ({type(e).__name__}: {e})")
        finally:
            doc.close()

    if merged.page_count == 0:
        merged.close()
        return None, warnings, encrypted

    out = io.BytesIO()
    merged.save(out, garbage=4, deflate=True, clean=True)
    merged.close()
    out.seek(0)
    return out.getvalue(), warnings, encrypted

# =========================================================
# SIGN DETECTION (text + OCR)
# =========================================================
def find_name_rect_text(doc: fitz.Document, target_text: str):
    for pi in range(doc.page_count):
        page = doc[pi]
        rects = page.search_for(target_text)
        if rects:
            return pi, rects[0]
    return None

def ocr_find_name_rect(doc: fitz.Document, target_text: str, zoom=2.8):
    if not OCR_AVAILABLE:
        return None
    target_text = (target_text or "").strip()
    if not target_text:
        return None
    target_tokens = normalize(target_text).split()
    if not target_tokens:
        return None

    for pi in range(doc.page_count):
        page = doc[pi]
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)

        words = []
        for i in range(len(data["text"])):
            txt = normalize(data["text"][i])
            if not txt:
                continue
            x, y, w, h = data["left"][i], data["top"][i], data["width"][i], data["height"][i]
            words.append((txt, x, y, x + w, y + h))

        for start in range(0, len(words) - len(target_tokens) + 1):
            ok = True
            for j, tok in enumerate(target_tokens):
                if words[start + j][0] != tok:
                    ok = False
                    break
            if ok:
                x0 = min(words[start + j][1] for j in range(len(target_tokens)))
                y0 = min(words[start + j][2] for j in range(len(target_tokens)))
                x1 = max(words[start + j][3] for j in range(len(target_tokens)))
                y1 = max(words[start + j][4] for j in range(len(target_tokens)))

                page_rect = page.rect
                sx = page_rect.width / pix.width
                sy = page_rect.height / pix.height
                rect_pdf = fitz.Rect(x0 * sx, y0 * sy, x1 * sx, y1 * sy)
                return pi, rect_pdf

    return None

def insert_signature_above_into_pdf(
    doc: fitz.Document,
    page_index: int,
    name_rect: fitz.Rect,
    sig_bytes: bytes,
    gap=6,
    pad=4,
    scale_w=1.4,
    scale_h=2.0,
    dx_pdf=0.0,
    dy_pdf=0.0,
):
    page = doc[page_index]
    name_w = name_rect.x1 - name_rect.x0
    name_h = name_rect.y1 - name_rect.y0

    w = name_w * scale_w + pad * 2
    h = name_h * scale_h + pad * 2

    cx = (name_rect.x0 + name_rect.x1) / 2
    x0 = cx - w / 2
    x1 = cx + w / 2

    y1 = name_rect.y0 - gap
    y0 = y1 - h

    x0 += dx_pdf
    x1 += dx_pdf
    y0 += dy_pdf
    y1 += dy_pdf

    if y0 < 0:
        y0 = 0
        y1 = h

    rect_sig = fitz.Rect(x0, y0, x1, y1)
    page.insert_image(rect_sig, stream=sig_bytes, overlay=True)

# =========================================================
# IMAGES → PDF (enhance)
# =========================================================
def open_uploaded_image(uploaded_image) -> Image.Image:
    img = Image.open(io.BytesIO(uploaded_image.getvalue()))
    img = ImageOps.exif_transpose(img)
    return img

def preprocess_image(img: Image.Image, auto_enhance=False,
                     brightness=1.0, contrast=1.0, sharpness=1.0,
                     grayscale=False, black_white=False) -> Image.Image:
    img = img.copy()
    if img.mode not in ("RGB", "RGBA", "L"):
        img = img.convert("RGB")

    if auto_enhance:
        temp = img.convert("RGB") if img.mode == "RGBA" else img.copy()
        temp = ImageOps.autocontrast(temp)
        temp = ImageEnhance.Contrast(temp).enhance(1.2)
        temp = ImageEnhance.Sharpness(temp).enhance(1.3)
        img = temp

    work = img.convert("RGB") if img.mode == "RGBA" else img.copy()
    work = ImageEnhance.Brightness(work).enhance(brightness)
    work = ImageEnhance.Contrast(work).enhance(contrast)
    work = ImageEnhance.Sharpness(work).enhance(sharpness)

    if grayscale:
        work = ImageOps.grayscale(work)
    if black_white:
        gray = ImageOps.grayscale(work)
        work = gray.point(lambda x: 255 if x > 160 else 0, mode="1").convert("RGB")
    else:
        if work.mode != "RGB":
            work = work.convert("RGB")
    return work

def pil_image_to_pdf_bytes(img: Image.Image) -> bytes:
    img = img.convert("RGB")
    out = io.BytesIO()
    img.save(out, format="PDF", resolution=150.0)
    out.seek(0)
    return out.getvalue()

def convert_multiple_images_to_individual_pdfs(uploaded_images, auto_enhance=False,
                                               brightness=1.0, contrast=1.0, sharpness=1.0,
                                               grayscale=False, black_white=False):
    pdfs, names = [], []
    for img_file in uploaded_images:
        img = open_uploaded_image(img_file)
        img = preprocess_image(img, auto_enhance, brightness, contrast, sharpness, grayscale, black_white)
        pdf_bytes = pil_image_to_pdf_bytes(img)
        base_name = img_file.name.rsplit(".", 1)[0]
        names.append(f"{base_name}.pdf")
        pdfs.append(pdf_bytes)
    return pdfs, names

def build_zip_of_pdfs(pdf_bytes_list, pdf_names):
    out_zip = io.BytesIO()
    with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
        for b, n in zip(pdf_bytes_list, pdf_names):
            zf.writestr(n, b)
    out_zip.seek(0)
    return out_zip.getvalue()

# =========================================================
# WORD (.DOCX) → PDF (sin dependencias del sistema)
# =========================================================
def _iter_docx_blocks(document):
    """Recorre párrafos y tablas respetando el orden del documento."""
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, document)

def _paragraph_runs(paragraph):
    """Incluye también el texto que está dentro de hipervínculos."""
    if hasattr(paragraph, "iter_inner_content"):
        for item in paragraph.iter_inner_content():
            if hasattr(item, "runs"):
                yield from item.runs
            else:
                yield item
    else:
        yield from paragraph.runs

def _paragraph_markup(paragraph):
    chunks = []
    for run in _paragraph_runs(paragraph):
        text = escape(run.text or "").replace("\t", "&nbsp;&nbsp;&nbsp;&nbsp;").replace("\n", "<br/>")
        if not text:
            continue
        if run.underline:
            text = f"<u>{text}</u>"
        if run.italic:
            text = f"<i>{text}</i>"
        if run.bold:
            text = f"<b>{text}</b>"
        chunks.append(text)
    return "".join(chunks)

def _paragraph_style(paragraph, styles, sequence):
    source_style = paragraph.style
    style_name = (source_style.name if source_style else "Normal").lower()
    if "title" in style_name or "título" in style_name:
        base = styles["Title"]
    elif "heading 1" in style_name or "título 1" in style_name:
        base = styles["Heading1"]
    elif "heading 2" in style_name or "título 2" in style_name:
        base = styles["Heading2"]
    elif "heading" in style_name or "título" in style_name:
        base = styles["Heading3"]
    else:
        base = styles["BodyText"]

    alignment = {
        0: TA_LEFT,
        1: TA_CENTER,
        2: TA_RIGHT,
        3: TA_JUSTIFY,
    }.get(paragraph.alignment, base.alignment)
    fmt = paragraph.paragraph_format
    font_size = None
    for run in paragraph.runs:
        if run.font.size:
            font_size = run.font.size.pt
            break
    if font_size is None and source_style and source_style.font.size:
        font_size = source_style.font.size.pt

    return ParagraphStyle(
        f"docx_style_{sequence}",
        parent=base,
        fontName="Helvetica",
        fontSize=font_size or base.fontSize,
        leading=(font_size or base.fontSize) * 1.25,
        alignment=alignment,
        leftIndent=fmt.left_indent.pt if fmt.left_indent else base.leftIndent,
        rightIndent=fmt.right_indent.pt if fmt.right_indent else base.rightIndent,
        firstLineIndent=fmt.first_line_indent.pt if fmt.first_line_indent else base.firstLineIndent,
        spaceBefore=fmt.space_before.pt if fmt.space_before else base.spaceBefore,
        spaceAfter=fmt.space_after.pt if fmt.space_after else max(base.spaceAfter, 4),
    )

def _paragraph_has_numbering(paragraph):
    properties = paragraph._p.pPr
    return properties is not None and properties.numPr is not None

def _paragraph_images(paragraph, document, available_width):
    images = []
    for blip in paragraph._p.xpath(".//a:blip"):
        relationship_id = blip.get(qn("r:embed"))
        if not relationship_id or relationship_id not in document.part.related_parts:
            continue
        try:
            image_bytes = document.part.related_parts[relationship_id].blob
            width, height = 4 * inch, 3 * inch
            extent = blip.xpath("ancestor::wp:inline/wp:extent | ancestor::wp:anchor/wp:extent")
            if extent:
                width = int(extent[0].get("cx")) / 914400 * inch
                height = int(extent[0].get("cy")) / 914400 * inch
            scale = min(1.0, available_width / width, (8.5 * inch) / height)
            image = ReportLabImage(io.BytesIO(image_bytes), width=width * scale, height=height * scale)
            image.hAlign = "CENTER"
            images.extend([image, Spacer(1, 6)])
        except Exception:
            # Una imagen dañada no debe impedir convertir el resto del documento.
            continue
    return images

def _table_flowable(docx_table, available_width, styles):
    column_count = max((len(row.cells) for row in docx_table.rows), default=1)
    data = []
    for row in docx_table.rows:
        cells = []
        for cell in row.cells:
            paragraphs = [_paragraph_markup(p) for p in cell.paragraphs]
            cell_markup = "<br/>".join(text for text in paragraphs if text) or " "
            cells.append(Paragraph(cell_markup, styles["BodyText"]))
        cells.extend([Paragraph(" ", styles["BodyText"])] * (column_count - len(cells)))
        data.append(cells)
    if not data:
        return Spacer(1, 1)
    table = Table(data, colWidths=[available_width / column_count] * column_count, repeatRows=1)
    table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B7B7B7")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F2F2")),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table

def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convierte un DOCX a PDF usando solo Python, apt/Word/LibreOffice no son necesarios."""
    document = Document(io.BytesIO(docx_bytes))
    section = document.sections[0]
    page_size = (section.page_width.pt, section.page_height.pt)
    left_margin = section.left_margin.pt
    right_margin = section.right_margin.pt
    top_margin = section.top_margin.pt
    bottom_margin = section.bottom_margin.pt
    available_width = page_size[0] - left_margin - right_margin

    output = io.BytesIO()
    pdf = SimpleDocTemplate(
        output,
        pagesize=page_size,
        leftMargin=left_margin,
        rightMargin=right_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin,
        title=document.core_properties.title or "Documento convertido",
        author=document.core_properties.author or "",
    )
    styles = getSampleStyleSheet()
    styles["BodyText"].fontName = "Helvetica"
    styles["BodyText"].fontSize = 10
    styles["BodyText"].leading = 12.5
    story = []

    for sequence, block in enumerate(_iter_docx_blocks(document)):
        if isinstance(block, DocxParagraph):
            markup = _paragraph_markup(block)
            style = _paragraph_style(block, styles, sequence)
            images = _paragraph_images(block, document, available_width)
            if markup:
                bullet = "•" if _paragraph_has_numbering(block) else None
                story.append(Paragraph(markup, style, bulletText=bullet))
            elif not images:
                story.append(Spacer(1, max(style.leading / 2, 4)))
            story.extend(images)
            if block._p.xpath(".//w:br[@w:type='page']"):
                story.append(PageBreak())
        else:
            story.extend([_table_flowable(block, available_width, styles), Spacer(1, 8)])

    if not story:
        story.append(Paragraph("Documento sin contenido visible.", styles["BodyText"]))

    header_text = " | ".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
    footer_text = " | ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())

    def draw_header_footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#666666"))
        if header_text:
            canvas.drawCentredString(page_size[0] / 2, page_size[1] - max(top_margin / 2, 18), header_text[:180])
        footer = footer_text[:150]
        if footer:
            footer = f"{footer}  ·  "
        canvas.drawCentredString(page_size[0] / 2, max(bottom_margin / 2, 14), f"{footer}Página {canvas.getPageNumber()}")
        canvas.restoreState()

    pdf.build(story, onFirstPage=draw_header_footer, onLaterPages=draw_header_footer)
    output.seek(0)
    return output.getvalue()

# =========================================================
# COMPRESS PDF
# =========================================================
def guess_file_type(filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".jpg") or name.endswith(".jpeg") or name.endswith(".png"):
        return "image"
    return "unknown"

def compress_pdf_soft(pdf_bytes: bytes) -> bytes:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    out = io.BytesIO()
    doc.save(out, garbage=4, deflate=True, clean=True)
    doc.close()
    out.seek(0)
    return out.getvalue()

def compress_pdf_rasterize(pdf_bytes: bytes, dpi: int = 120, jpeg_quality: int = 60) -> bytes:
    src = fitz.open(stream=pdf_bytes, filetype="pdf")
    dst = fitz.open()
    zoom = dpi / 72.0

    for i in range(src.page_count):
        page = src[i]
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        img_bytes = io.BytesIO()
        img.save(img_bytes, format="JPEG", quality=jpeg_quality, optimize=True)

        rect = page.rect
        new_page = dst.new_page(width=rect.width, height=rect.height)
        new_page.insert_image(rect, stream=img_bytes.getvalue())

    out = io.BytesIO()
    dst.save(out, garbage=4, deflate=True, clean=True)
    dst.close()
    src.close()
    out.seek(0)
    return out.getvalue()

# =========================================================
# SPLIT PDF HELPERS
# =========================================================
def parse_ranges(range_text: str, max_pages: int):
    """
    Texto: '1-3,5,7-10' -> [(1,3),(5,5),(7,10)] (1-based)
    """
    range_text = (range_text or "").strip().replace(" ", "")
    if not range_text:
        return []
    parts = range_text.split(",")
    ranges = []
    for p in parts:
        if not p:
            continue
        if "-" in p:
            a, b = p.split("-", 1)
            a = int(a); b = int(b)
        else:
            a = int(p); b = int(p)
        if a < 1: a = 1
        if b > max_pages: b = max_pages
        if a > b:
            a, b = b, a
        ranges.append((a, b))
    return ranges

def zip_outputs(files):
    zbio = io.BytesIO()
    with zipfile.ZipFile(zbio, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, b in files:
            zf.writestr(name, b)
    zbio.seek(0)
    return zbio.getvalue()

def split_doc_each_page(doc: fitz.Document):
    outputs = []
    for i in range(doc.page_count):
        out_doc = fitz.open()
        out_doc.insert_pdf(doc, from_page=i, to_page=i, links=0)
        bio = io.BytesIO()
        out_doc.save(bio, garbage=4, deflate=True, clean=True)
        out_doc.close()
        outputs.append((f"pagina_{i+1:03d}.pdf", bio.getvalue()))
    return outputs

def split_doc_by_ranges(doc: fitz.Document, ranges_1based):
    outputs = []
    for idx, (a, b) in enumerate(ranges_1based, start=1):
        out_doc = fitz.open()
        out_doc.insert_pdf(doc, from_page=a-1, to_page=b-1, links=0)
        bio = io.BytesIO()
        out_doc.save(bio, garbage=4, deflate=True, clean=True)
        out_doc.close()
        outputs.append((f"rango_{idx:02d}_{a:03d}-{b:03d}.pdf", bio.getvalue()))
    return outputs

def split_doc_every_n(doc: fitz.Document, n: int):
    n = max(1, int(n))
    outputs = []
    total = doc.page_count
    part = 1
    for start in range(0, total, n):
        end = min(total - 1, start + n - 1)
        out_doc = fitz.open()
        out_doc.insert_pdf(doc, from_page=start, to_page=end, links=0)
        bio = io.BytesIO()
        out_doc.save(bio, garbage=4, deflate=True, clean=True)
        out_doc.close()
        outputs.append((f"parte_{part:02d}_{start+1:03d}-{end+1:03d}.pdf", bio.getvalue()))
        part += 1
    return outputs

# =========================================================
# SIDEBAR
# =========================================================
st.sidebar.title("📚 Menú")
menu = st.sidebar.radio(
    "Selecciona una herramienta",
    ["Inicio", "Unir PDFs", "Dividir PDF", "Word a PDF", "Imágenes a PDF", "Comprimir PDF"],
)
st.sidebar.markdown("---")
st.sidebar.caption("Carga → previsualiza (opcional) → procesa → descarga ✅")

# =========================================================
# HOME
# =========================================================
if menu == "Inicio":
    st.markdown('<div class="hero">Aplicativo en construcción para Karina 💓</div>', unsafe_allow_html=True)
    st.markdown(
        f"<div class='card'><h2 style='margin:0'>{APP_TITLE}</h2>"
        f"<p class='muted'>Rápido, con previsualización bajo demanda y manejo de PDFs con contraseña.</p></div>",
        unsafe_allow_html=True
    )

# =========================================================
# MODULE: UNIR PDFs
# =========================================================
elif menu == "Unir PDFs":
    st.markdown(
        "<div class='card'>"
        "<h2 style='margin:0'>📄 Unir PDFs</h2>"
        "<p class='muted'>Carga → (preview opcional) → ordena → une → (preview opcional) → firma opcional → descarga.</p>"
        "</div>",
        unsafe_allow_html=True
    )

    top1, top2, top3 = st.columns([1, 1, 2])
    with top1:
        if st.button("🔄 Reiniciar módulo"):
            reset_merge()
            st.rerun()
    with top2:
        enable_ocr = st.toggle(
            "Usar OCR si viene escaneado",
            value=False,
            disabled=not OCR_AVAILABLE,
            help="Requiere Tesseract instalado en el servidor.",
            key="enable_ocr",
        )
    with top3:
        st.session_state.last_target = st.text_input(
            "Nombre a detectar (para habilitar firma)",
            value=st.session_state.last_target,
            key="target_name"
        )

    st.session_state.last_output_name = st.text_input(
        "Nombre del PDF final",
        value=st.session_state.last_output_name,
        key="output_name"
    )

    uploaded_files = st.file_uploader(
        "1) Carga tus PDFs",
        type=["pdf"],
        accept_multiple_files=True,
        key="pdfs_uploader"
    )

    # Sync inteligente por conjunto
    if uploaded_files:
        incoming_items = [(f.name, len(f.getvalue())) for f in uploaded_files]
        incoming_signature = tuple(sorted(incoming_items))
        if st.session_state.merge_signature != incoming_signature:
            st.session_state.merge_files = [
                {"name": f.name, "bytes": f.getvalue(), "size": len(f.getvalue())} for f in uploaded_files
            ]
            st.session_state.merge_signature = incoming_signature

    if uploaded_files and st.button("↩️ Sincronizar con selección del explorador"):
        st.session_state.merge_files = [{"name": f.name, "bytes": f.getvalue(), "size": len(f.getvalue())} for f in uploaded_files]
        st.rerun()

    # Preview rápido de un archivo
    if st.session_state.merge_files:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.subheader("2) Vista rápida de un archivo (opcional)")
        idx = st.selectbox(
            "Selecciona un archivo",
            options=list(range(len(st.session_state.merge_files))),
            format_func=lambda i: st.session_state.merge_files[i]["name"],
            key="merge_preview_selected_file"
        )
        preview_first_page_quick(
            st.session_state.merge_files[idx]["bytes"],
            key_prefix="merge_file_quick",
            label=f"Primera página: {st.session_state.merge_files[idx]['name']}"
        )
        st.markdown("</div>", unsafe_allow_html=True)

    # Ordenamiento
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.subheader("3) Ordena los archivos (orden final)")

    if st.session_state.merge_files:
        a1, a2, a3, a4 = st.columns([1, 1, 1, 1])
        with a1:
            if st.button("A→Z"):
                st.session_state.merge_files = sorted(st.session_state.merge_files, key=lambda x: x["name"].lower())
                st.rerun()
        with a2:
            if st.button("🔁 Invertir"):
                st.session_state.merge_files = list(reversed(st.session_state.merge_files))
                st.rerun()
        with a3:
            if st.button("🧹 Limpiar lista"):
                st.session_state.merge_files = []
                st.session_state.merge_signature = None
                st.rerun()
        with a4:
            st.caption(f"{len(st.session_state.merge_files)} archivo(s)")

        for i, item in enumerate(st.session_state.merge_files):
            c1, c2, c3, c4 = st.columns([7, 1, 1, 1])
            with c1:
                mb = item["size"] / (1024 * 1024)
                st.write(f"**{i+1}.** {item['name']} · {mb:.2f} MB")
            with c2:
                if st.button("⬆️", key=f"merge_up_{i}") and i > 0:
                    st.session_state.merge_files[i-1], st.session_state.merge_files[i] = (
                        st.session_state.merge_files[i],
                        st.session_state.merge_files[i-1],
                    )
                    st.rerun()
            with c3:
                if st.button("⬇️", key=f"merge_down_{i}") and i < len(st.session_state.merge_files) - 1:
                    st.session_state.merge_files[i+1], st.session_state.merge_files[i] = (
                        st.session_state.merge_files[i],
                        st.session_state.merge_files[i+1],
                    )
                    st.rerun()
            with c4:
                if st.button("🗑️", key=f"merge_del_{i}"):
                    st.session_state.merge_files.pop(i)
                    st.rerun()
    else:
        st.info("Carga PDFs para crear la cola.")

    st.markdown("</div>", unsafe_allow_html=True)

    # Password UI
    if st.session_state.encrypted_files:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.error("🔒 Hay PDFs con contraseña. Ingresa la clave y vuelve a presionar **Unir PDFs**.")

        common_pwd = st.text_input("Contraseña (si es la misma para todos)", type="password", key="merge_common_pwd")
        cA, cB = st.columns([1, 1])
        with cA:
            if st.button("Aplicar a todos"):
                for fn in st.session_state.encrypted_files:
                    st.session_state.pdf_passwords[fn] = common_pwd
                st.success("Aplicado ✅")
        with cB:
            if st.button("🧹 Limpiar contraseñas"):
                st.session_state.pdf_passwords = {}
                st.success("Listo ✅")

        for fn in st.session_state.encrypted_files:
            key = f"merge_pwd_{fn}"
            st.text_input(f"Contraseña para: {fn}", type="password",
                          value=st.session_state.pdf_passwords.get(fn, ""), key=key)
            st.session_state.pdf_passwords[fn] = st.session_state[key]

        st.markdown("</div>", unsafe_allow_html=True)

    # Merge action
    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.subheader("4) Unir")

    can_merge = len(st.session_state.merge_files) >= 1
    if st.button("✅ Unir PDFs", disabled=not can_merge, key="do_merge"):
        st.session_state.encrypted_files = []

        with st.spinner("Uniendo PDFs..."):
            merged_bytes, warnings, encrypted = merge_pdfs_with_password(st.session_state.merge_files)

        if warnings:
            st.warning("Algunos archivos tuvieron problemas:")
            for w in warnings:
                st.write(w)

        if encrypted:
            st.session_state.encrypted_files = encrypted
            st.session_state.merged_pdf_bytes = None
            st.error("Hay PDFs protegidos. Ingresa contraseñas y vuelve a intentar.")
        else:
            st.session_state.merged_pdf_bytes = merged_bytes
            st.success("✅ PDF unido listo. Previsualiza y descarga.")

            # Detect name (safe)
            target = (st.session_state.last_target or "").strip()
            found, method = None, None
            if target:
                doc = fitz.open(stream=merged_bytes, filetype="pdf")
                found = find_name_rect_text(doc, target)
                method = "texto"
                if (not found) and enable_ocr:
                    method = "ocr"
                    with st.spinner("Buscando por OCR..."):
                        found_ocr = ocr_find_name_rect(doc, target, zoom=2.8)
                    if found_ocr:
                        found = found_ocr
                doc.close()

            if found:
                page_index, rect = found
                st.session_state.detected = True
                st.session_state.det_page = int(page_index)
                st.session_state.det_rect = (float(rect.x0), float(rect.y0), float(rect.x1), float(rect.y1))
                st.session_state.det_method = method
            else:
                st.session_state.detected = False
                st.session_state.det_page = None
                st.session_state.det_rect = None
                st.session_state.det_method = None

    st.markdown("</div>", unsafe_allow_html=True)

    # Result preview + download + signing
    if st.session_state.merged_pdf_bytes:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.subheader("5) Resultado")

        preview_pdf_viewer_fast(
            st.session_state.merged_pdf_bytes,
            key_prefix="merge_result_preview",
            title="Previsualización del PDF unido"
        )

        outname = st.session_state.last_output_name
        outname = outname if outname.lower().endswith(".pdf") else outname + ".pdf"
        st.download_button(
            "⬇️ Descargar PDF unido",
            data=st.session_state.merged_pdf_bytes,
            file_name=outname,
            mime="application/pdf",
            key="dl_merge_out"
        )

        target = (st.session_state.last_target or "").strip()
        if not target:
            st.info("💡 Si quieres firma, escribe el **Nombre a detectar**.")
        elif not st.session_state.detected:
            st.info("No se detectó el nombre en el PDF unido.")
        else:
            st.success(f"Nombre detectado ✅ Método: {st.session_state.det_method} | Página: {st.session_state.det_page + 1}")

            wants_sign = st.toggle("6) Firmar (opcional)", value=False, key="wants_sign")
            if wants_sign:
                sig_file = st.file_uploader("Sube la firma (PNG/JPG)", type=["png", "jpg", "jpeg"], key="sig_uploader")

                gap = st.slider("Espacio entre firma y nombre", 0, 40, 10, key="gap")
                pad = st.slider("Margen", 0, 20, 6, key="pad")
                scale_w = st.slider("Escala ancho firma", 0.8, 2.5, 1.4, 0.1, key="scale_w")
                scale_h = st.slider("Escala alto firma", 0.8, 4.0, 2.0, 0.1, key="scale_h")

                step = st.slider("Paso movimiento", 1, 30, 6, key="move_step")
                c1, c2, c3, c4, c5 = st.columns(5)
                with c1:
                    if st.button("⬅️", key="btn_left"):
                        st.session_state.sig_dx -= float(step)
                with c2:
                    if st.button("➡️", key="btn_right"):
                        st.session_state.sig_dx += float(step)
                with c3:
                    if st.button("⬆️", key="btn_up"):
                        st.session_state.sig_dy -= float(step)
                with c4:
                    if st.button("⬇️", key="btn_down"):
                        st.session_state.sig_dy += float(step)
                with c5:
                    if st.button("Reset", key="btn_reset"):
                        st.session_state.sig_dx = 0.0
                        st.session_state.sig_dy = 0.0

                if sig_file and st.button("🔒 Generar PDF firmado", type="primary", key="make_signed"):
                    doc2 = fitz.open(stream=st.session_state.merged_pdf_bytes, filetype="pdf")
                    rect_pdf = fitz.Rect(*st.session_state.det_rect)

                    insert_signature_above_into_pdf(
                        doc2,
                        st.session_state.det_page,
                        rect_pdf,
                        sig_file.getvalue(),
                        gap=6,
                        pad=pad,
                        scale_w=scale_w,
                        scale_h=scale_h,
                        dx_pdf=st.session_state.sig_dx,
                        dy_pdf=st.session_state.sig_dy,
                    )

                    out = io.BytesIO()
                    doc2.save(out)
                    doc2.close()
                    out.seek(0)
                    signed_bytes = out.getvalue()

                    st.markdown("---")
                    preview_pdf_viewer_fast(
                        signed_bytes,
                        key_prefix="signed_preview",
                        title="Previsualización del PDF firmado"
                    )
                    st.download_button(
                        "⬇️ Descargar PDF firmado",
                        data=signed_bytes,
                        file_name="PDF_unido_firmado.pdf",
                        mime="application/pdf",
                        key="dl_signed"
                    )

        st.markdown("</div>", unsafe_allow_html=True)

# =========================================================
# MODULE: DIVIDIR PDF
# =========================================================
elif menu == "Dividir PDF":
    st.markdown(
        "<div class='card'>"
        "<h2 style='margin:0'>✂️ Dividir PDF</h2>"
        "<p class='muted'>Carga → previsualiza (opcional) → elige método → descarga ZIP.</p>"
        "</div>",
        unsafe_allow_html=True
    )

    c1, c2 = st.columns([1, 3])
    with c1:
        if st.button("🔄 Reiniciar módulo"):
            reset_split()
            st.rerun()

    split_file = st.file_uploader("Sube el PDF a dividir", type=["pdf"], key="split_uploader")

    if split_file:
        pdf_bytes = split_file.getvalue()
        fname = split_file.name

        # Intentar abrir (para saber si está encriptado / páginas)
        doc, err = open_pdf_with_password(pdf_bytes, fname)

        if err == "encrypted":
            st.session_state.encrypted_split = [fname]
            st.error("🔒 Este PDF está protegido. Ingresa la contraseña para continuar.")

            pwd_key = f"split_pwd_{fname}"
            st.text_input("Contraseña", type="password", value=st.session_state.pdf_passwords.get(fname, ""), key=pwd_key)
            st.session_state.pdf_passwords[fname] = st.session_state[pwd_key]

            if st.button("🔓 Validar contraseña"):
                doc2, err2 = open_pdf_with_password(pdf_bytes, fname)
                if err2 == "encrypted":
                    st.error("Contraseña incorrecta.")
                elif doc2 is None:
                    st.error(f"No se pudo abrir: {err2}")
                else:
                    st.success("Contraseña OK ✅")
                    doc2.close()

        elif doc is None:
            st.error(f"No se pudo abrir el PDF: {err}")
        else:
            total_pages = doc.page_count
            doc.close()

            preview_pdf_viewer_fast(pdf_bytes, key_prefix="split_preview", title="Previsualización del PDF cargado")
            st.caption(f"Total páginas: {total_pages}")

            metodo = st.radio(
                "¿Cómo deseas dividir?",
                [
                    "Cada página por separado (ZIP)",
                    "Por rangos (ej: 1-3,5,7-10)",
                    "En partes de N páginas (ej: cada 10)"
                ],
                key="split_method"
            )

            outputs = None

            if metodo == "Cada página por separado (ZIP)":
                if st.button("✂️ Dividir ahora", type="primary", key="split_go_1"):
                    with st.spinner("Dividiendo..."):
                        doc3, err3 = open_pdf_with_password(pdf_bytes, fname)
                        if doc3 is None:
                            st.error(f"No se pudo abrir: {err3}")
                        else:
                            outputs = split_doc_each_page(doc3)
                            doc3.close()

            elif metodo == "Por rangos (ej: 1-3,5,7-10)":
                range_text = st.text_input("Rangos", value="1-1", key="split_ranges")
                if st.button("✂️ Dividir ahora", type="primary", key="split_go_2"):
                    with st.spinner("Dividiendo..."):
                        ranges = parse_ranges(range_text, total_pages)
                        if not ranges:
                            st.error("Rangos inválidos. Ejemplo: 1-3,5,7-10")
                        else:
                            doc3, err3 = open_pdf_with_password(pdf_bytes, fname)
                            if doc3 is None:
                                st.error(f"No se pudo abrir: {err3}")
                            else:
                                outputs = split_doc_by_ranges(doc3, ranges)
                                doc3.close()

            else:  # partes de N
                n = st.number_input("Páginas por parte", min_value=1, value=10, step=1, key="split_n")
                if st.button("✂️ Dividir ahora", type="primary", key="split_go_3"):
                    with st.spinner("Dividiendo..."):
                        doc3, err3 = open_pdf_with_password(pdf_bytes, fname)
                        if doc3 is None:
                            st.error(f"No se pudo abrir: {err3}")
                        else:
                            outputs = split_doc_every_n(doc3, int(n))
                            doc3.close()

            if outputs:
                st.session_state.split_outputs = outputs
                st.success(f"Listo ✅ Se generaron {len(outputs)} archivo(s).")

    # Descargas y preview de una salida
    if st.session_state.split_outputs:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.subheader("⬇️ Descarga")

        z = zip_outputs(st.session_state.split_outputs)
        st.download_button(
            "Descargar ZIP con PDFs divididos",
            data=z,
            file_name="pdf_dividido.zip",
            mime="application/zip",
            key="dl_split_zip"
        )

        # Preview rápida de una salida (opcional)
        st.markdown("### 👀 Vista rápida de una salida (opcional)")
        out_idx = st.selectbox(
            "Selecciona un archivo generado",
            options=list(range(len(st.session_state.split_outputs))),
            format_func=lambda i: st.session_state.split_outputs[i][0],
            key="split_out_select"
        )
        preview_first_page_quick(
            st.session_state.split_outputs[out_idx][1],
            key_prefix="split_out_preview",
            label=f"Primera página: {st.session_state.split_outputs[out_idx][0]}"
        )

        with st.expander("Descargas individuales (muestra hasta 30)"):
            for fname, fbytes in st.session_state.split_outputs[:30]:
                st.download_button(
                    f"Descargar {fname}",
                    data=fbytes,
                    file_name=fname,
                    mime="application/pdf",
                    key=f"dl_{fname}"
                )
            if len(st.session_state.split_outputs) > 30:
                st.info("Mostrando 30. Usa el ZIP para descargar todo.")

        st.markdown("</div>", unsafe_allow_html=True)

# =========================================================
# MODULE: WORD A PDF
# =========================================================
elif menu == "Word a PDF":
    st.markdown(
        "<div class='card'><h2 style='margin:0'>📝 Word a PDF</h2>"
        "<p class='muted'>Convierte documentos DOCX a PDF sin enviar el archivo a servicios externos.</p></div>",
        unsafe_allow_html=True
    )

    if st.button("🔄 Reiniciar conversión", key="reset_word"):
        reset_word()
        st.rerun()

    word_file = st.file_uploader(
        "Sube el documento de Word",
        type=["docx"],
        help="Se admite el formato moderno .docx. Los archivos antiguos .doc deben guardarse primero como .docx.",
        key="word_uploader",
    )

    if word_file:
        if st.session_state.word_source_name != word_file.name:
            st.session_state.word_pdf_bytes = None
            st.session_state.word_source_name = word_file.name
        default_name = f"{word_file.name.rsplit('.', 1)[0]}.pdf"
        output_name = st.text_input(
            "Nombre del PDF",
            value=default_name,
            key=f"word_output_{word_file.name}",
        )
        st.caption("Se conservan textos, estilos comunes, listas, tablas, imágenes, márgenes y saltos de página.")

        if st.button("📝 Convertir Word a PDF", type="primary", key="convert_word"):
            try:
                with st.spinner("Convirtiendo documento..."):
                    converted = convert_docx_to_pdf(word_file.getvalue())
                st.session_state.word_pdf_bytes = converted
                st.session_state.word_pdf_name = output_name if output_name.lower().endswith(".pdf") else output_name + ".pdf"
                st.success("✅ Documento convertido correctamente.")
            except Exception as exc:
                st.session_state.word_pdf_bytes = None
                st.error(f"No se pudo convertir el DOCX: {type(exc).__name__}: {exc}")

    if st.session_state.word_pdf_bytes:
        preview_pdf_viewer_fast(
            st.session_state.word_pdf_bytes,
            key_prefix="word_preview",
            title="Previsualización del documento convertido",
        )
        st.download_button(
            "⬇️ Descargar PDF",
            data=st.session_state.word_pdf_bytes,
            file_name=st.session_state.word_pdf_name,
            mime="application/pdf",
            key="download_word_pdf",
        )

# =========================================================
# MODULE: IMÁGENES A PDF
# =========================================================
elif menu == "Imágenes a PDF":
    st.markdown(
        "<div class='card'><h2 style='margin:0'>🖼️ Imágenes a PDF</h2>"
        "<p class='muted'>Convierte imágenes a PDF con mejora opcional. Preview del unificado bajo demanda.</p></div>",
        unsafe_allow_html=True
    )

    if st.button("🔄 Reiniciar sección imágenes"):
        reset_images()
        st.rerun()

    uploaded_images = st.file_uploader(
        "Sube una o varias imágenes",
        type=["jpg", "jpeg", "png"],
        accept_multiple_files=True,
        key="images_uploader"
    )

    if uploaded_images:
        selected_preview = st.selectbox(
            "Imagen para previsualizar",
            options=list(range(len(uploaded_images))),
            format_func=lambda i: uploaded_images[i].name,
            key="selected_preview_img"
        )
        original_img = open_uploaded_image(uploaded_images[selected_preview])

        auto_enhance = st.toggle("Mejora automática", value=True, key="auto_enhance")
        c1, c2, c3 = st.columns(3)
        with c1:
            brightness = st.slider("Brillo", 0.5, 2.0, 1.0, 0.1, key="brightness")
        with c2:
            contrast = st.slider("Contraste", 0.5, 2.5, 1.2, 0.1, key="contrast")
        with c3:
            sharpness = st.slider("Nitidez", 0.5, 3.0, 1.2, 0.1, key="sharpness")
        c4, c5 = st.columns(2)
        with c4:
            grayscale = st.checkbox("Escala de grises", key="grayscale")
        with c5:
            black_white = st.checkbox("Blanco y negro", key="black_white")

        improved_img = preprocess_image(original_img, auto_enhance, brightness, contrast, sharpness, grayscale, black_white)

        p1, p2 = st.columns(2)
        with p1:
            st.image(original_img, caption="Original", use_container_width=True)
        with p2:
            st.image(improved_img, caption="Mejorada", use_container_width=True)

        mode = st.radio(
            "¿Qué deseas hacer?",
            ["Convertir y descargar PDFs individuales", "Convertir y unificar en un solo PDF", "Hacer ambas opciones"],
            key="images_mode"
        )

        if st.button("🖼️ Convertir imágenes", key="convert_images_btn"):
            with st.spinner("Convirtiendo..."):
                pdfs, names = convert_multiple_images_to_individual_pdfs(
                    uploaded_images, auto_enhance, brightness, contrast, sharpness, grayscale, black_white
                )

            st.session_state.converted_images_pdf_bytes = pdfs
            st.session_state.converted_images_names = names

            if mode in ["Convertir y unificar en un solo PDF", "Hacer ambas opciones"]:
                merged = fitz.open()
                for b in pdfs:
                    src = fitz.open(stream=b, filetype="pdf")
                    merged.insert_pdf(src, links=0)
                    src.close()
                out = io.BytesIO()
                merged.save(out, garbage=4, deflate=True, clean=True)
                merged.close()
                out.seek(0)
                st.session_state.merged_images_pdf_bytes = out.getvalue()
            else:
                st.session_state.merged_images_pdf_bytes = None

            st.success("✅ Listo. Descarga disponible.")

    if st.session_state.converted_images_pdf_bytes:
        st.markdown("<div class='card'>", unsafe_allow_html=True)
        st.subheader("⬇️ Descargas")

        mode = st.session_state.get("images_mode", "Convertir y descargar PDFs individuales")

        if mode in ["Convertir y descargar PDFs individuales", "Hacer ambas opciones"]:
            zip_bytes = build_zip_of_pdfs(
                st.session_state.converted_images_pdf_bytes,
                st.session_state.converted_images_names
            )
            st.download_button(
                "Descargar ZIP (PDFs individuales)",
                data=zip_bytes,
                file_name="imagenes_convertidas_pdf.zip",
                mime="application/zip",
                key="download_zip_individuals"
            )

        if mode in ["Convertir y unificar en un solo PDF", "Hacer ambas opciones"] and st.session_state.merged_images_pdf_bytes:
            preview_pdf_viewer_fast(
                st.session_state.merged_images_pdf_bytes,
                key_prefix="images_merged_preview",
                title="Previsualización del PDF unificado"
            )
            st.download_button(
                "Descargar PDF unificado",
                data=st.session_state.merged_images_pdf_bytes,
                file_name="imagenes_unificadas.pdf",
                mime="application/pdf",
                key="download_merged_images_pdf"
            )

        st.markdown("</div>", unsafe_allow_html=True)

# =========================================================
# MODULE: COMPRIMIR PDF
# =========================================================
elif menu == "Comprimir PDF":
    st.markdown(
        "<div class='card'><h2 style='margin:0'>🗜️ Comprimir PDF</h2>"
        "<p class='muted'>Modo rápido reduce mucho (ideal escaneados). Preview bajo demanda.</p></div>",
        unsafe_allow_html=True
    )

    if st.button("🔄 Reiniciar compresión"):
        reset_compress()
        st.rerun()

    up = st.file_uploader("Sube tu archivo", type=["pdf", "jpg", "jpeg", "png"], key="compress_uploader")
    if up:
        ftype = guess_file_type(up.name)
        st.info(f"Tipo detectado: {ftype.upper()}")

        if ftype != "pdf":
            st.warning("Este módulo comprime PDFs. Si subiste imagen, usa Imágenes a PDF.")
        else:
            original_bytes = up.getvalue()
            original_mb = len(original_bytes) / (1024 * 1024)
            st.write(f"Tamaño original: **{original_mb:.2f} MB**")

            mode = st.radio(
                "Modo",
                ["Rápido (reduce mucho)", "Suave (reduce poco, conserva mejor)"],
                index=0,
                key="compress_mode"
            )

            if mode.startswith("Rápido"):
                dpi = st.select_slider("Calidad (DPI)", options=[72, 96, 120, 150], value=120, key="dpi")
                quality = st.select_slider("Calidad JPEG", options=[35, 45, 60, 75], value=60, key="jpg_quality")
            else:
                dpi, quality = None, None

            out_name = st.text_input("Nombre del PDF comprimido", value="archivo_comprimido.pdf", key="compress_outname")

            if st.button("🗜️ Comprimir ahora", type="primary", key="do_compress"):
                with st.spinner("Comprimiendo..."):
                    if mode.startswith("Rápido"):
                        compressed = compress_pdf_rasterize(original_bytes, dpi=int(dpi), jpeg_quality=int(quality))
                    else:
                        compressed = compress_pdf_soft(original_bytes)

                st.session_state.compressed_pdf_bytes = compressed
                st.session_state.compressed_name = out_name if out_name.lower().endswith(".pdf") else out_name + ".pdf"

                new_mb = len(compressed) / (1024 * 1024)
                reduction = (1 - (new_mb / original_mb)) * 100 if original_mb > 0 else 0
                st.success(f"Listo ✅ Nuevo tamaño: **{new_mb:.2f} MB** | Reducción aprox: **{reduction:.1f}%**")

    if st.session_state.compressed_pdf_bytes:
        preview_pdf_viewer_fast(
            st.session_state.compressed_pdf_bytes,
            key_prefix="compress_preview",
            title="Previsualización del PDF comprimido"
        )
        st.download_button(
            "Descargar PDF comprimido",
            data=st.session_state.compressed_pdf_bytes,
            file_name=st.session_state.compressed_name,
            mime="application/pdf",
            key="download_compressed"
        )
